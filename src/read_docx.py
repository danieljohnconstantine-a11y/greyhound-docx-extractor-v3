from docx import Document

def _table_to_matrix(tbl):
    matrix = []
    for row in tbl.rows:
        matrix.append([cell.text.strip() if cell.text else "" for cell in row.cells])
    return matrix

def load_docx(file_path: str):
    """
    Load a .docx file and extract all visible content:
    - paragraphs
    - tables (as matrices)
    - headers / footers (all sections)
    Returns a dictionary with raw content blocks for parsing.
    """
    doc = Document(file_path)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip() != ""]

    tables = []
    for t in doc.tables:
        tables.append(_table_to_matrix(t))

    headers = []
    footers = []
    try:
        for sec in doc.sections:
            if sec.header:
                for p in sec.header.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        headers.append(txt)
            if sec.footer:
                for p in sec.footer.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        footers.append(txt)
    except Exception:
        # Some documents may not expose header/footer cleanly; ignore if not present
        pass

    raw_text = "\n".join(headers + paragraphs + footers)

    return {
        "file_path": file_path,
        "headers": headers,
        "footers": footers,
        "paragraphs": paragraphs,
        "tables": tables,
        "raw_text": raw_text,
    }
