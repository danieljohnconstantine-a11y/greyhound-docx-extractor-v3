import re
from datetime import datetime

import pandas as pd
from docx import Document

from src.summary_utils import normalize_summary_fields
from src.columns import HISTORY_COLUMNS


def read_docx_tables(path: str):
    doc = Document(path)
    blocks = []

    # paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            blocks.append({"type": "paragraph", "text": t})

    # tables (in order they appear)
    for tbl in doc.tables:
        rows = []
        for r in tbl.rows:
            cells = [c.text.strip() for c in r.cells]
            if any(x for x in cells):
                rows.append(cells)
        if rows:
            blocks.append({"type": "table", "rows": rows})

    return blocks


def parse_meeting_header(blocks, source_file: str):
    info = {
        "Track": "",
        "Race_Date": "",
        "Race_No": "",
        "Distance_m": "",
        "Race_Name": "",
        "Race_Grade": "",
        "Data_Source_File": source_file,
    }

    for b in blocks:
        if b["type"] != "paragraph":
            continue

        t = b["text"]

        # Race number
        m = re.search(r"Race\s*No\.?\s*(\d+)", t, re.I)
        if m:
            info["Race_No"] = m.group(1)

        # Date (e.g. 15 Oct 25)
        m = re.search(r"(\d{1,2}\s*[A-Za-z]{3}\s*\d{2,4})", t)
        if m:
            info["Race_Date"] = m.group(1)

        # Track name
        m = re.search(
            r"(ALBANY|ANGLE PARK|BALLARAT|BENDIGO|CANBERRA|CANNINGTON|CAPALABA|DUBBO|DAWSON|GRAFTON|GOSFORD|HEALESVILLE|HORSHAM|IPSWICH|LAUNCESTON|MANDURAH|MURRAY BRIDGE|NORTHAM|RICHMOND|SALE|SANDOWN|TAREE|THE MEADOWS|TOWNSVILLE|TRARALGON|WAGGA|WARRAGUL|WARRNAMBOOL|WENTWORTH PARK|YOUNG)",
            t,
            re.I,
        )
        if m:
            info["Track"] = m.group(1).upper()

        # Distance (e.g. 520m)
        m = re.search(r"(\d{3,4})\s*m", t, re.I)
        if m:
            info["Distance_m"] = m.group(1)

        # Race grade
        m = re.search(r"Grade\s*([A-Za-z0-9/ -]+)", t, re.I)
        if m:
            info["Race_Grade"] = m.group(1).strip()

        # Race name fallback (first long-ish line if not yet set)
        if not info["Race_Name"]:
            if " " in t and len(t.split()) >= 3:
                info["Race_Name"] = t

    if info["Race_Date"]:
        try:
            info["Race_Date"] = pd.to_datetime(info["Race_Date"], dayfirst=True).strftime(
                "%Y-%m-%d"
            )
        except Exception:
            pass

    return info


def map_header(h: str) -> str:
    h = h.upper()

    # dog identity
    if "DOG" in h:
        return "Dog_Name"
    if "BOX" in h:
        return "Box"
    if "TAB" in h:
        return "Tab_No"
    if "TRAINER" in h:
        return "Trainer"
    if "WT" in h or "WEIGHT" in h:
        return "Weight_kg"
    if "FORM" in h:
        return "FF_Form"
    if h == "BP" or "BOX NO" in h:
        return "BP"

    # pedigree
    if "SIRE" in h:
        return "Sire"
    if "DAM" in h:
        return "Dam"
    if "OWNER" in h:
        return "Owner"

    # age/sex
    if "AGE" in h:
        return "Age"
    if "SEX" in h or "GENDER" in h:
        return "Sex"

    # career stats
    if "PRIZE" in h:
        return "PrizeMoney"
    if "WINS" in h:
        return "Wins"
    if "PLACES" in h:
        return "Places"
    if "STARTS" in h:
        return "Starts"
    if "DLR" in h:
        return "DLR"
    if "DLW" in h:
        return "DLW"
    if "RTC" in h:
        return "RTC"

    return h.title()


def parse_dog_table_rows(table_rows, meeting_info):
    rows = []
    headers = None

    for r in table_rows:
        if not headers:
            headers = [h.strip().upper() for h in r]
            continue

        row_dict = dict(meeting_info)

        for idx, val in enumerate(r):
            h = headers[idx] if idx < len(headers) else ""
            key = map_header(h)
            row_dict[key] = val.strip()

        if "Dog_Name" not in row_dict:
            if "DOG" in row_dict:
                row_dict["Dog_Name"] = row_dict["DOG"]

        if "Box" not in row_dict:
            if "BOX" in row_dict:
                row_dict["Box"] = row_dict["BOX"]

        if "Trainer" not in row_dict:
            if "TRAINER" in row_dict:
                row_dict["Trainer"] = row_dict["TRAINER"]

        rows.append(row_dict)

    return rows


def extract_dog_tables(blocks, meeting_info):
    all_rows = []
    for b in blocks:
        if b["type"] != "table":
            continue

        table_rows = b["rows"]
        if len(table_rows) < 2:
            continue

        header_row = [c.upper() for c in table_rows[0]]
        if any(x in header_row for x in ["DOG", "BOX", "TAB", "TRAINER"]):
            parsed = parse_dog_table_rows(table_rows, meeting_info)
            all_rows.extend(parsed)

    return all_rows


def map_history_header(h: str) -> str | None:
    h = h.upper()
    if "DATE" in h:
        return "Hist_Date"
    if "TRACK" in h or "VENUE" in h:
        return "Hist_Track"
    if "DIST" in h or "DIS" in h:
        return "Hist_Distance"
    if "POS" in h or "PLC" in h or "PLACE" in h:
        return "Hist_Finish_Pos"
    if "MARGIN" in h or "MARG" in h:
        return "Hist_Margin_L"
    if "TIME" in h and "SEC" not in h:
        return "Hist_Race_Time"
    if "SEC" in h and "ADJ" not in h:
        return "Hist_Sec_Time"
    if "ADJ" in h:
        return "Hist_Sec_Time_Adj"
    if "SOT" in h or "TRK" in h:
        return "Hist_SOT"
    if "RST" in h:
        return "Hist_RST"
    if "BOX" in h or "BP" in h:
        return "Hist_BP"
    if "ODDS" in h or "SP" in h:
        return "Hist_Odds"
    if "API" in h:
        return "Hist_API"
    if "PRIZE" in h or "STAKE" in h or "STK" in h:
        return "Hist_Prize_Won"
    if "WINNER" in h:
        return "Hist_Winner"
    if "2ND" in h:
        return "Hist_2nd_Place"
    if "3RD" in h:
        return "Hist_3rd_Place"
    if "SETTLE" in h or "SETTL" in h:
        return "Hist_Settled_Turn"
    if "ONGOING" in h:
        return "Hist_Ongoing_Winners"
    if "DIR" in h or "RAIL" in h:
        return "Hist_Track_Direction"
    return None


def _to_float_or_none(v):
    if v is None:
        return None
    s = str(v).strip()
    if s == "":
        return None
    try:
        return float(s)
    except Exception:
        return None


def _parse_time_to_seconds(s: str | None) -> float | None:
    if s is None:
        return None
    text = str(s).strip()
    if not text:
        return None

    # typical race time "30.12"
    m = re.match(r"^(\d{2,3}\.\d{1,3})$", text)
    if m:
        try:
            return float(m.group(1))
        except Exception:
            return None

    # allow plain float
    try:
        return float(text)
    except Exception:
        return None


def _compute_speed_kmh(dist_str: str | None, time_str: str | None) -> float | None:
    dist_m = None
    if dist_str is not None:
        m = re.search(r"(\d{2,4})", str(dist_str))
        if m:
            try:
                dist_m = float(m.group(1))
            except Exception:
                dist_m = None

    t_s = _parse_time_to_seconds(time_str)

    if dist_m is None or t_s is None or t_s == 0:
        return None

    return (dist_m * 3.6) / t_s


def parse_history_blocks(blocks, meeting_info, dog_rows):
    """
    Parse history tables into list of dicts with keys in HISTORY_COLUMNS.
    Links tables to dogs by scanning paragraphs that mention known dog names.
    """

    # Known dog names (for context linking)
    dog_names = sorted(
        {r.get("Dog_Name", "").strip() for r in dog_rows if r.get("Dog_Name")},
        key=len,
        reverse=True,
    )
    dog_box = {}
    for r in dog_rows:
        name = (r.get("Dog_Name") or "").strip()
        if name:
            dog_box[name] = (r.get("Box") or "").strip()

    hist_rows = []
    current_dog = None

    for b in blocks:
        if b["type"] == "paragraph":
            text = b["text"]
            if not dog_names:
                continue
            for name in dog_names:
                # whole-word search
                if re.search(r"\b" + re.escape(name) + r"\b", text):
                    current_dog = name
                    break

        elif b["type"] == "table":
            rows = b["rows"]
            if len(rows) < 2:
                continue

            header = [c.upper() for c in rows[0]]
            # detect history tables: require DATE + (DIST or MARGIN or TIME)
            if "DATE" not in " ".join(header):
                continue
            if not any(x in " ".join(header) for x in ["DIST", "DIS", "MARGIN", "TIME"]):
                continue

            # map headers to history fields
            col_map = {}
            for idx, h in enumerate(header):
                key = map_history_header(h)
                if key:
                    col_map[idx] = key

            # if we can't map at least Date + Distance or Date + Time, skip
            has_date = any(v == "Hist_Date" for v in col_map.values())
            has_distance_or_time = any(
                v in ["Hist_Distance", "Hist_Race_Time"] for v in col_map.values()
            )
            if not (has_date and has_distance_or_time):
                continue

            # parse each data row
            for r in rows[1:]:
                hist = {k: "" for k in HISTORY_COLUMNS}

                hist["Track"] = meeting_info.get("Track", "")
                hist["Race_Date"] = meeting_info.get("Race_Date", "")
                hist["Race_No"] = meeting_info.get("Race_No", "")
                hist["Dog_Name"] = current_dog or ""
                hist["Box"] = dog_box.get(current_dog or "", "")
                hist["Data_Source_File"] = meeting_info.get("Data_Source_File", "")

                for idx, val in enumerate(r):
                    if idx not in col_map:
                        continue
                    key = col_map[idx]
                    hist[key] = val.strip()

                # compute per-run speed
                spd = _compute_speed_kmh(hist.get("Hist_Distance"), hist.get("Hist_Race_Time"))
                if spd is not None:
                    hist["Hist_Speed_km/h"] = round(spd, 3)
                else:
                    hist["Hist_Speed_km/h"] = ""

                hist_rows.append(hist)

    return hist_rows


def parse_docx(path: str):
    blocks = read_docx_tables(path)
    meeting_info = parse_meeting_header(blocks, source_file=path)

    dog_rows = extract_dog_tables(blocks, meeting_info)
    hist_rows = parse_history_blocks(blocks, meeting_info, dog_rows)

    df = pd.DataFrame(dog_rows)
    df = normalize_summary_fields(df)

    return df, hist_rows
